import random
import streamlit as st
import gen
with st.sidebar:
    api_key = st.text_input("你的api key(默认为空即可)", key="file_qa_api_key", type="password")
    "[狠狠赞助我](https://afdian.net/a/voidfun/plan)"

st.title("📝 vv o r d --- 一个Ai辅助背单词助手")
st.subheader("step1 -你要背的单词")
words = st.text_input("空格隔开即可")
option = st.selectbox(
    '或者选择单词书',
    ('','高中'))
if option =="高中":
    file = open("gz.txt","r").readlines()
    for i in range(5):
        words += file[random.randint(0,1139)]
    st.write(words)
    pass
st.subheader("step2 -选择故事风格(可以多选)")
options = st.multiselect(
    '故事风格',
    ['浪漫', '悬疑', '恐怖', '穿越','古代','幽默','黑色','感动','反转','武侠','童话','科幻','幻想'])
style:str=""
if len(options) != 0:
    for i in options:
        if len(style) == 0:
            style=i
        else:
            style+=','+i
genable=False
if words and style:
    genable=True
st.subheader("step2.5 -可选项")
ex = st.checkbox("生成word文档")
if not ex:
    tl = st.checkbox("听力训练(必须勾选生成文档)",disabled=True)
else:
    tl = st.checkbox("听力训练(没实现)", disabled=True)
lm = st.checkbox("长故事")
st.checkbox("故事视频",disabled=True)
st.subheader("step3 -AI赋予你力量")
if st.button("生成故事",disabled=not genable):
    try:
        with st.spinner("生成文章中 -请稍候"):
            genable = False
            if api_key:
                TEXT = gen.gen(words, option, longmethod=lm,api=True)
            else:
                TEXT = gen.gen(words, option, longmethod=lm)
            genable = True
        if ex:
            with st.spinner("整合文档中"):
                from docx import Document
                path = f"doc/{words}.docx"
                doc = Document(path)
                doc.add_heading(words, level=1)
                paragraph1 = doc.add_paragraph(TEXT)
                doc.save(path)
    except Exception as e:
        st.write(e)
        st.error("what did u do")
        st.rerun()
    st.balloons()
    st.info("生成文章成功")
    if ex:
        st.download_button("立即下载",file_name=f"doc/{words}.docx",mime="application/vnd.openxmlformats-officedocument");
    else:
        st.write(TEXT)
